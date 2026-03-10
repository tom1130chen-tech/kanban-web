#!/usr/bin/env python3
"""
Document Formatter - ENGN 2125 Class Notes
美化排版学习笔记文档
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def format_document(input_path, output_path):
    """美化文档排版"""
    
    # 读取原文档
    doc = Document(input_path)
    
    # 创建新文档
    new_doc = Document()
    
    # 设置页面
    section = new_doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    
    # 定义颜色
    BLUE = RGBColor(41, 98, 255)
    DARK_BLUE = RGBColor(23, 55, 94)
    GRAY = RGBColor(128, 128, 128)
    LIGHT_GRAY = RGBColor(240, 240, 240)
    RED = RGBColor(231, 76, 60)
    
    # 收集所有内容
    content = []
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text.strip())
    
    # 添加标题
    title = new_doc.add_heading('ENGN 2125 Class 07', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = DARK_BLUE
    
    subtitle = new_doc.add_paragraph('管理决策心理学 · 学习笔记')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.name = 'Calibri'
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = GRAY
    
    new_doc.add_paragraph()  # 空行
    
    # 处理内容
    i = 0
    while i < len(content):
        line = content[i]
        
        # 一级标题（一、二、三...）
        if line.startswith(('一、', '二、', '三、', '四、', '五、')):
            heading = new_doc.add_heading(line, level=1)
            heading.runs[0].font.name = 'Calibri'
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.bold = True
            heading.runs[0].font.color.rgb = BLUE
            i += 1
        
        # 二级标题（1. 2. 3...）
        elif line[0].isdigit() and '. ' in line[:3]:
            heading = new_doc.add_heading(line, level=2)
            heading.runs[0].font.name = 'Calibri'
            heading.runs[0].font.size = Pt(14)
            heading.runs[0].font.bold = True
            heading.runs[0].font.color.rgb = DARK_BLUE
            i += 1
        
        # 表格标记 - 跳过
        elif line == 'Table' or line == 'plain' or line == 'Copy':
            i += 1
        
        # 核心洞察/关键案例等特殊段落
        elif line.startswith(('核心洞察：', '核心概念：', '核心框架', '关键案例：', '核心问题：', '核心管理启示')):
            p = new_doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.right_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('💡 ' + line)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = BLUE
            # 添加浅蓝色背景（通过 shading 实现）
            i += 1
        
        # 定义/机制
        elif line.startswith(('定义：', '机制：')):
            p = new_doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(line)
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
            i += 1
        
        # 典型表现/经典案例
        elif line.startswith(('典型表现：', '经典案例：', '管理应用：', '商业应用：', '实验证据', '干预方法：', '成因：', '影响：', '表现：', '关键案例：', '例外情况')):
            p = new_doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run('▸ ' + line)
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = DARK_BLUE
            i += 1
        
        # 列表项（破折号或特殊符号）
        elif line.startswith(('—', '•', '◦')):
            p = new_doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(line)
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            i += 1
        
        # 普通内容
        else:
            p = new_doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(50, 50, 50)
            i += 1
    
    # 添加页脚信息
    new_doc.add_paragraph()
    footer = new_doc.add_paragraph('ENGN 2125 · 布朗大学工程学院')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.name = 'Calibri'
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = GRAY
    footer.runs[0].font.italic = True
    
    # 保存
    new_doc.save(output_path)
    print(f"✅ 文档已保存：{output_path}")
    return output_path

if __name__ == '__main__':
    input_path = '/Users/tomchen/.openclaw/media/inbound/c9da561c-debc-4cd4-bf9d-1a8ef6179c7d.docx'
    output_path = '/Users/tomchen/.openclaw/workspace-chat/kanban-board/temp/ENGN2125_Class07_Notes_Formatted.docx'
    format_document(input_path, output_path)
